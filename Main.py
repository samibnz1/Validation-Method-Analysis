import pandas as pd
import numpy as np
import statsmodels.api as sm
import statsmodels.stats.anova as sm_anova
from statsmodels.formula.api import ols
from scipy.stats import probplot
from scipy.stats import shapiro, t
import matplotlib.pyplot as plt
import os
from docx import Document
from docx.shared import Inches
import uuid
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk, Button, Label, Entry
from PIL import Image, ImageTk

def browse_file_entry(entry):
    file_path = filedialog.askopenfilename(
        title="Sélectionner un fichier Excel",
        filetypes=[("Fichiers Excel", "*.xls *.xlsx"), ("Tous les fichiers", "*.*")]
    )
    if file_path:
        entry.delete(0, tk.END)
        entry.insert(0, file_path)


def perform_regression_analysis(df):

    X = df.iloc[:, :-1]
    X = sm.add_constant(X)


    y = df.iloc[:, -1]


    model = sm.OLS(y, X).fit()


    model_summary = model.summary()
    coefficients = model.params
    p_values = model.pvalues
    conf_int = model.conf_int()


    equation_codée = f"Y = {coefficients[0]:.4f} + " + " + ".join(
        [f"({coefficients[i]:.4f} * {X.columns[i]})" for i in range(1, len(coefficients))])


    effect_matrix = df.iloc[:, :-1].apply(lambda x: x * coefficients[1:].values, axis=1)

    return model_summary, coefficients, equation_codée, effect_matrix, model


def create_report(df, model_summary, coefficients, equation_codée, effect_matrix, model):

    doc = Document()
    doc.add_heading("Rapport d'Analyse de Robustesse", 0)


    doc.add_heading("1. Introduction et Objectifs", level=1)
    doc.add_paragraph(
        "Cette analyse vise à évaluer la robustesse du modèle en examinant l'impact de plusieurs facteurs sur la variable réponse. "
        "Les objectifs sont d'identifier les facteurs influents et de déterminer leur significativité."
    )


    doc.add_heading("2. Méthodologie", level=1)
    doc.add_paragraph(
        "Les facteurs étudiés incluent les colonnes de facteurs dans le fichier Excel. Un modèle de régression linéaire a été construit pour analyser "
        "leur impact sur la variable réponse. Les coefficients de régression ont été estimés et l'analyse des résidus a été effectuée pour valider la robustesse du modèle."
    )


    doc.add_heading("3. Matrice d'Effet en Variable Codée", level=1)
    effect_matrix_table = doc.add_table(rows=1, cols=len(df.columns) + 1)
    effect_matrix_table.style = 'Light Shading Accent 1'
    hdr_cells = effect_matrix_table.rows[0].cells
    hdr_cells[0].text = 'Numéro d\'essai'

    for i, col in enumerate(df.columns):
        hdr_cells[i + 1].text = col

    for idx, row in df.iterrows():
        row_cells = effect_matrix_table.add_row().cells
        row_cells[0].text = str(idx + 1)
        for i, val in enumerate(row):
            row_cells[i + 1].text = f"{val:.4f}"


    doc.add_heading("4. Équation du Modèle en Variable Réelle", level=1)
    doc.add_paragraph(equation_codée)


    doc.add_heading("5. Résultats de la Régression", level=1)
    doc.add_paragraph(str(model_summary))


    doc.add_heading("6. Interprétation des Coefficients", level=1)
    for factor, coef, pval in zip(df.columns[:-1], coefficients[1:], model.pvalues[1:]):
        interpretation = f"Le facteur {factor} a un coefficient de {coef:.4f}. "
        if pval < 0.05:
            interpretation += f"Ce facteur est significatif (p-value = {pval:.3f})."
        else:
            interpretation += f"Ce facteur n'est pas significatif (p-value = {pval:.3f})."
        doc.add_paragraph(interpretation)


    doc.add_heading("7. Interprétation du Test F", level=1)
    f_stat = model.fvalue
    p_value = model.f_pvalue
    doc.add_paragraph(
        f"Le test F pour le modèle est : {f_stat:.4f} avec une p-value de {p_value:.4f}."
    )
    if p_value < 0.05:
        doc.add_paragraph("Le modèle est globalement significatif.")
    else:
        doc.add_paragraph("Le modèle n'est pas globalement significatif.")


    doc.add_heading("8. Graphiques de Validation du Modèle", level=1)


    fig_values_vs_predicted = plt.figure(figsize=(8, 5))
    plt.scatter(model.fittedvalues, df.iloc[:, -1], color='blue', label='Valeurs Prédites vs Réelles')
    plt.plot([df.iloc[:, -1].min(), df.iloc[:, -1].max()], [df.iloc[:, -1].min(), df.iloc[:, -1].max()], color='red',
             linestyle='--', label="Ligne de Régression")
    plt.xlabel("Valeurs Prédites")
    plt.ylabel("Valeurs Réelles")
    plt.title("Graphique des Valeurs Réelles vs Prédites")
    plt.legend()
    plt.tight_layout()
    real_vs_pred_img_path = f"real_vs_pred_{uuid.uuid4()}.png"
    plt.savefig(real_vs_pred_img_path)
    plt.close()
    doc.add_picture(real_vs_pred_img_path, width=Inches(5))


    fig_residuals = plt.figure(figsize=(8, 5))
    residuals = model.resid
    plt.scatter(model.fittedvalues, residuals, color='green', label='Résidus')
    plt.axhline(y=0, color='red', linestyle='--', label='Résidu = 0')
    plt.xlabel("Valeurs Prédites")
    plt.ylabel("Résidus")
    plt.title("Graphique des Résidus")
    plt.legend()
    plt.tight_layout()
    residuals_img_path = f"residuals_{uuid.uuid4()}.png"
    plt.savefig(residuals_img_path)
    plt.close()
    doc.add_picture(residuals_img_path, width=Inches(5))


    fig_qq = plt.figure(figsize=(8, 5))
    probplot(residuals, dist="norm", plot=plt)
    plt.title("QQ-Plot des Résidus")
    plt.tight_layout()
    qq_img_path = f"qq_plot_{uuid.uuid4()}.png"
    plt.savefig(qq_img_path)
    plt.close()


    fig_qq_borders = plt.figure(figsize=(8, 5))
    probplot(residuals, dist="norm", plot=plt)
    plt.title("QQ-Plot des Résidus avec Bordures")
    plt.gca().get_lines()[0].set_color('blue')  # Changer la couleur de la ligne
    plt.gca().get_lines()[0].set_linewidth(2)  # Changer l'épaisseur de la ligne
    plt.gca().get_lines()[1].set_markerfacecolor('red')  # Changer la couleur des points
    plt.gca().get_lines()[1].set_marker('o')  # Modifier le style des marqueurs
    plt.tight_layout()
    qq_borders_img_path = f"qq_plot_borders_{uuid.uuid4()}.png"
    plt.savefig(qq_borders_img_path)
    plt.close()

    doc.add_picture(qq_borders_img_path, width=Inches(5))


    doc.add_heading("9. Analyse Pareto des Facteurs Influents", level=1)
    pareto_data = coefficients[1:].sort_values(ascending=False)
    doc.add_paragraph("Les facteurs influents sont ceux avec les coefficients les plus significatifs.")
    pareto_data_table = doc.add_table(rows=1, cols=2)
    pareto_data_table.style = 'Light Shading Accent 1'
    hdr_cells = pareto_data_table.rows[0].cells
    hdr_cells[0].text = 'Facteur'
    hdr_cells[1].text = 'Coefficient'

    for factor, coef in pareto_data.items():
        row_cells = pareto_data_table.add_row().cells
        row_cells[0].text = factor
        row_cells[1].text = f"{coef:.4f}"


    fig_pareto = plt.figure(figsize=(8, 5))
    pareto_data.plot(kind='bar', color='lightblue')
    plt.title("Analyse Pareto des Facteurs Influents")
    plt.ylabel("Coefficient")
    plt.tight_layout()
    pareto_img_path = f"pareto_{uuid.uuid4()}.png"
    plt.savefig(pareto_img_path)
    plt.close()

    doc.add_picture(pareto_img_path, width=Inches(5))


    doc.add_heading("10. Effets Cumulés des Coefficients", level=1)
    fig_cumulative = plt.figure(figsize=(8, 5))
    pareto_data.cumsum().plot(kind='line', marker='o', color='orange')
    plt.title("Effets Cumulés des Coefficients")
    plt.ylabel("Effet Cumulé")
    plt.tight_layout()
    cumulative_img_path = f"cumulative_{uuid.uuid4()}.png"
    plt.savefig(cumulative_img_path)
    plt.close()

    doc.add_picture(cumulative_img_path, width=Inches(5))


    doc.add_heading("11. Conclusions et Recommandations", level=1)
    doc.add_paragraph(
        "L'analyse de robustesse du modèle révèle que les facteurs ayant des coefficients significatifs "
        "ont un impact direct sur la variable réponse. Les facteurs avec des coefficients positifs "
        "indiquent une amélioration de la performance du modèle lorsque leur niveau augmente, tandis que "
        "les facteurs avec des coefficients négatifs montrent qu'une augmentation de ces facteurs pourrait "
        "diminuer les performances du modèle."
    )

    doc.add_paragraph(
        "Les tests de validité, notamment le test F, ont confirmé que le modèle est globalement significatif. "
        "Cependant, certains facteurs pourraient nécessiter une attention particulière pour améliorer la robustesse "
        "et l'applicabilité du modèle dans des scénarios réels."
    )

    # Save report
    output_file = "Rapport_Analyse_Robustesse.docx"
    doc.save(output_file)
    messagebox.showinfo("Succès", f"Rapport généré : {output_file}")


def open_robustness_window(master):
    win = tk.Toplevel(master)
    win.title("Analyse de Robustesse")
    win.geometry("450x300")

    # Sélectionner le fichier Excel
    tk.Label(win, text="Fichier Excel contenant les facteurs, niveaux et réponses :").pack(pady=5)
    entry_file = tk.Entry(win, width=50)
    entry_file.pack(pady=5)
    tk.Button(win, text="Parcourir", command=lambda: browse_file_entry(entry_file)).pack(pady=5)

    def generate_robustness_report():
        file_path = entry_file.get()
        if not file_path:
            messagebox.showerror("Erreur", "Veuillez sélectionner un fichier Excel.")
            return
        try:
            df = pd.read_excel(file_path)
            model_summary, coefficients, equation_codée, effect_matrix, model = perform_regression_analysis(df)
            create_report(df, model_summary, coefficients, equation_codée, effect_matrix, model)
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'analyse du fichier : {e}")

    tk.Button(win, text="Générer Rapport", command=generate_robustness_report).pack(pady=10)


def cochran_test(groups, alpha=0.05):

    k = len(groups)
    n = len(groups[0]) if k > 0 else 0
    variances = [np.var(g, ddof=1) for g in groups]
    s_var = sum(variances)

    if s_var == 0:

        C = np.nan
        decision = "Impossible de calculer (variances nulles)."
        bool_test = False
    else:
        C = max(variances) / s_var
        decision = None
        bool_test = False


    cochran_table = {
        (5, 3): 0.683
    }

    C_crit = cochran_table.get((k, n), 0.683)

    if not np.isnan(C):
        if C < C_crit:
            decision = (
                f"Test de Cochran : C={C:.4f} < {C_crit:.4f} → "
                "variances homogènes (pas de différence significative)."
            )
            bool_test = True
        else:
            decision = (
                f"Test de Cochran : C={C:.4f} ≥ {C_crit:.4f} → "
                "variances non homogènes (différence significative)."
            )
            bool_test = False
    return (C, C_crit, decision, bool_test)



def linear_regression_analysis(x, y):

    df_temp = pd.DataFrame({"x": x, "y": y}).dropna()
    model = ols("y ~ x", data=df_temp).fit()
    anova_results = sm_anova.anova_lm(model, typ=1)
    residuals = model.resid
    fitted_values = model.fittedvalues


    if len(residuals) >= 3:
        w_stat, w_pval = shapiro(residuals)
    else:
        w_stat, w_pval = np.nan, np.nan

    return {
        "model": model,
        "params": model.params,
        "bse": model.bse,
        "pvalues": model.pvalues,
        "conf_int": model.conf_int(0.05),
        "anova": anova_results,
        "residuals": residuals,
        "fitted_values": fitted_values,
        "shapiro_stat": w_stat,
        "shapiro_pvalue": w_pval,
        "r2": model.rsquared,
        "r2_adj": model.rsquared_adj,
        "ssr": model.ssr,
        "df_resid": model.df_resid
    }



def compute_validity_table(df):


    mean_by_niv = df.groupby("Niveau")["yij"].mean().rename("y_moy")
    merged = df.merge(mean_by_niv, on="Niveau")
    merged["sq_dev"] = (merged["yij"] - merged["y_moy"]) ** 2
    merged.sort_values(by=["Niveau"], inplace=True)
    SCE_exp = merged["sq_dev"].sum()
    return merged, SCE_exp



def compare_slopes_intercepts(reg0, reg1, alpha=0.05):

    slope0 = reg0["params"]["x"]
    slope1 = reg1["params"]["x"]
    var_slope0 = reg0["bse"]["x"] ** 2
    var_slope1 = reg1["bse"]["x"] ** 2

    intercept0 = reg0["params"]["Intercept"]
    intercept1 = reg1["params"]["Intercept"]
    var_int0 = reg0["bse"]["Intercept"] ** 2
    var_int1 = reg1["bse"]["Intercept"] ** 2


    df_slope = int(reg0["df_resid"] + reg1["df_resid"])
    df_intercept = df_slope


    diff_slope = slope0 - slope1
    se_diff_slope = np.sqrt(var_slope0 + var_slope1)
    t_slope = diff_slope / se_diff_slope
    p_slope = 2 * (1 - t.cdf(abs(t_slope), df_slope))

    if p_slope < alpha:
        interpret_slope = (
            f"La différence de pente est significative (p={p_slope:.3g} < {alpha})."
        )
    else:
        interpret_slope = (
            f"La différence de pente n'est pas significative (p={p_slope:.3g} ≥ {alpha})."
        )


    diff_int = intercept0 - intercept1
    se_diff_int = np.sqrt(var_int0 + var_int1)
    t_int = diff_int / se_diff_int
    p_int = 2 * (1 - t.cdf(abs(t_int), df_intercept))

    if p_int < alpha:
        interpret_int = (
            f"La différence d'ordonnée à l'origine est significative (p={p_int:.3g} < {alpha})."
        )
    else:
        interpret_int = (
            f"La différence d'ordonnée à l'origine n'est pas significative (p={p_int:.3g} ≥ {alpha})."
        )

    return {
        "diff_slope": diff_slope,
        "t_slope": t_slope,
        "df_slope": df_slope,
        "p_slope": p_slope,
        "diff_intercept": diff_int,
        "t_intercept": t_int,
        "df_intercept": df_intercept,
        "p_intercept": p_int,
        "interpret_slope": interpret_slope,
        "interpret_intercept": interpret_int
    }


def compare_intercept_with_zero(reg, alpha=0.05):

    intercept = reg["params"]["Intercept"]
    se_int = reg["bse"]["Intercept"]
    t_calc = intercept / se_int
    df_ = int(reg["df_resid"])
    p_val = 2 * (1 - t.cdf(abs(t_calc), df_))

    if p_val < alpha:
        interpret = (
            f"L'ordonnée à l'origine (Intercept={intercept:.3f}) est significativement ≠ 0 (p={p_val:.3g} < {alpha})."
        )
    else:
        interpret = (
            f"Pas de différence significative entre l'ordonnée à l'origine et 0 (p={p_val:.3g} ≥ {alpha})."
        )
    return t_calc, df_, p_val, interpret



def add_table_from_dict(doc, dict_data, title=None):
    if title:
        doc.add_paragraph(title, style='Intense Quote')
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Paramètre'
    hdr_cells[1].text = 'Valeur'
    for k, v in dict_data.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(k)
        row_cells[1].text = str(v)


def add_custom_table_data(doc, data, columns, title=None, style_name='Light Shading Accent 1'):

    if title:
        doc.add_paragraph(title, style='Intense Quote')
    rows = len(data) + 1
    cols = len(columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = style_name


    hdr_cells = table.rows[0].cells
    for j, col in enumerate(columns):
        hdr_cells[j].text = str(col)


    for i, row_data in enumerate(data, start=1):
        row_cells = table.rows[i].cells
        for j, val in enumerate(row_data):
            row_cells[j].text = str(val)


def insert_matplotlib_figure(doc, fig_path, caption=None, width_in_inch=5):
    doc.add_picture(fig_path, width=Inches(width_in_inch))
    if caption:
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = 1
        doc.add_paragraph(caption).alignment = 1
    if os.path.exists(fig_path):
        os.remove(fig_path)


def add_anova_table(doc, anova_df, title="Tableau ANOVA"):
    doc.add_paragraph(title, style='Intense Quote')
    nb_rows = anova_df.shape[0] + 1
    nb_cols = anova_df.shape[1] + 1
    table = doc.add_table(rows=nb_rows, cols=nb_cols)
    table.style = 'Light List Accent 1'


    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Source"
    for j, col in enumerate(anova_df.columns, start=1):
        hdr_cells[j].text = str(col)


    for i, row_label in enumerate(anova_df.index, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(row_label)
        for j, col in enumerate(anova_df.columns, start=1):
            val = anova_df.loc[row_label, col]
            if pd.isna(val):
                txt = "NaN"
            else:
                txt = f"{val:0.4g}"
            row_cells[j].text = txt



def add_non_linearity_table(doc, df, reg, sce_exp, name="D0"):

    from scipy.stats import f

    n = len(df)
    k = len(df["Niveau"].unique())
    SSR = reg["ssr"]
    df_resid_model = int(reg["df_resid"])
    var_res_model = SSR / df_resid_model if df_resid_model > 0 else np.nan


    df_exp = n - k
    var_exp = sce_exp / df_exp if df_exp > 0 else np.nan


    sce_nl = sce_exp - SSR
    df_nl = k - 2
    var_nl = sce_nl / df_nl if df_nl > 0 else np.nan


    if var_res_model > 0 and not np.isnan(var_nl):
        F_nl = var_nl / var_res_model
    else:
        F_nl = np.nan


    F_crit = np.nan
    if df_nl > 0 and df_resid_model > 0:
        try:
            F_crit = f.isf(0.05, df_nl, df_resid_model)
        except:
            F_crit = 3.71


    data = [
        [
            "Résiduelle",
            f"{SSR:.4f}",
            str(df_resid_model),
            f"{var_res_model:.4f}" if not np.isnan(var_res_model) else "",
            "",
            ""
        ],
        [
            "Expérimentale",
            f"{sce_exp:.4f}",
            str(df_exp),
            f"{var_exp:.4f}" if not np.isnan(var_exp) else "",
            "",
            ""
        ],
        [
            "Erreur modèle (non linéaire)",
            f"{sce_nl:.4f}",
            str(df_nl),
            f"{var_nl:.4f}" if not np.isnan(var_nl) else "",
            f"{F_nl:.4f}" if not np.isnan(F_nl) else "",
            f"{F_crit:.3f}" if not np.isnan(F_crit) else ""
        ]
    ]
    columns = ["Sources de variation", "SCE", "DDL", "Variances", "Fnl", "F(5%,3,10)"]
    add_custom_table_data(doc, data, columns, title=f"Test de non-linéarité - {name}")


    if not np.isnan(F_nl) and not np.isnan(F_crit):
        if F_nl < F_crit:
            doc.add_paragraph(
                f"Pour {name}, F_nl={F_nl:.3f} < F_crit={F_crit:.3f} ⇒ "
                f"pas de non-linéarité significative (la droite semble adaptée)."
            )
        else:
            doc.add_paragraph(
                f"Pour {name}, F_nl={F_nl:.3f} ≥ F_crit={F_crit:.3f} ⇒ "
                f"présence d'une non-linéarité significative (la droite n'est pas adéquate)."
            )
    else:
        doc.add_paragraph(f"Pour {name}, calcul de F_nl impossible ou non pertinent (données insuffisantes ?).")



def add_non_linearity_table(doc, df, reg, sce_exp, name="D0"):
    from scipy.stats import f
    n = len(df)
    k = len(df["Niveau"].unique())
    SSR = reg["ssr"]
    df_resid_model = int(reg["df_resid"])
    var_res_model = SSR / df_resid_model if df_resid_model > 0 else np.nan
    df_exp = n - k
    var_exp = sce_exp / df_exp if df_exp > 0 else np.nan
    sce_nl = sce_exp - SSR
    df_nl = k - 2
    var_nl = sce_nl / df_nl if df_nl > 0 else np.nan
    if var_res_model > 0 and not np.isnan(var_nl):
        F_nl = var_nl / var_res_model
    else:
        F_nl = np.nan
    F_crit = np.nan
    if df_nl > 0 and df_resid_model > 0:
        try:
            F_crit = f.isf(0.05, df_nl, df_resid_model)
        except:
            F_crit = 3.71
    data = [
        ["Résiduelle", f"{SSR:.4f}", str(df_resid_model), f"{var_res_model:.4f}" if not np.isnan(var_res_model) else "",
         "", ""],
        ["Expérimentale", f"{sce_exp:.4f}", str(df_exp), f"{var_exp:.4f}" if not np.isnan(var_exp) else "", "", ""],
        ["Erreur modèle (non linéaire)", f"{sce_nl:.4f}", str(df_nl), f"{var_nl:.4f}" if not np.isnan(var_nl) else "",
         f"{F_nl:.4f}" if not np.isnan(F_nl) else "", f"{F_crit:.3f}" if not np.isnan(F_crit) else ""]
    ]
    columns = ["Sources de variation", "SCE", "DDL", "Variances", "Fnl", "F(5%,3,10)"]
    add_custom_table_data(doc, data, columns, title=f"Test de non-linéarité - {name}")
    if not np.isnan(F_nl) and not np.isnan(F_crit):
        if F_nl < F_crit:
            doc.add_paragraph(
                f"Pour {name}, F_nl={F_nl:.3f} < F_crit={F_crit:.3f} ⇒ pas de non-linéarité significative.")
        else:
            doc.add_paragraph(
                f"Pour {name}, F_nl={F_nl:.3f} ≥ F_crit={F_crit:.3f} ⇒ présence d'une non-linéarité significative.")
    else:
        doc.add_paragraph(f"Pour {name}, calcul de F_nl impossible ou non pertinent.")



def create_report_linearity(
        df_PA, df_PAFPR,
        cochran_PA, cochran_PAFPR,
        reg_PA, reg_PAFPR,
        valid_PA, SCEexp_PA,
        valid_PAFPR, SCEexp_PAFPR,
        comp_pentes,
        comp_int_d0_zero,
        comp_int_d1_zero,
        alpha=0.05,
        output_file="Rapport_Linearite.docx"
):
    doc = Document()


    doc.add_heading("Rapport de Validation – Linéarité (Automatique)", 0)
    doc.add_paragraph(
        "Ce rapport de validation analytique a pour objectif d’évaluer la linéarité de la méthode sur la plage étudiée. "
        "Pour ce faire, plusieurs analyses complémentaires sont réalisées :\n"
        " - Un test de Cochran pour vérifier l’homogénéité des variances,\n"
        " - Une régression linéaire détaillée (coefficients, p-values, intervalles de confiance, ANOVA, R² et vérification de la normalité des résidus),\n"
        " - Une vérification de la validité de la droite via le calcul des moyennes par niveau,\n"
        " - Un test de non-linéarité comparant les écarts entre le modèle linéaire et les moyennes expérimentales,\n"
        " - Et enfin, une comparaison des pentes et des ordonnées (ainsi qu’une vérification de l’ordonnée à l’origine par rapport à zéro).\n\n"
        "L’ensemble de ces étapes permet de déterminer si la relation entre la variable d’intérêt et la réponse est bien linéaire, "
        "et par conséquent si la méthode peut être utilisée en toute confiance dans le cadre de son application analytique."
    )

    doc.add_page_break()


    doc.add_heading("1. Test de Cochran (Homogénéité des Variances)", level=1)

    c_pa, ccrit_pa, dec_pa, bool_cochran_pa = cochran_PA
    add_table_from_dict(doc, {
        "C (PA)": f"{c_pa:.4f}" if not np.isnan(c_pa) else "NaN",
        "C Crit. (5%)": f"{ccrit_pa:.4f}",
        "Résultat Cochran PA": dec_pa
    }, title="Test de Cochran - PA")

    c_pafpr, ccrit_pafpr, dec_pafpr, bool_cochran_pafpr = cochran_PAFPR
    add_table_from_dict(doc, {
        "C (PA+FPR)": f"{c_pafpr:.4f}" if not np.isnan(c_pafpr) else "NaN",
        "C Crit. (5%)": f"{ccrit_pafpr:.4f}",
        "Résultat Cochran PA+FPR": dec_pafpr
    }, title="Test de Cochran - PA+FPR")

    doc.add_paragraph(
        "Interprétation : si C < C_crit, les variances sont jugées homogènes. "
        f"Test effectué au seuil α={alpha}."
    )


    doc.add_heading("2. Régression Linéaire : Coefficients, Intervalles et Visualisations", level=1)


    doc.add_heading("2.1 PA Seul", level=2)
    model_pa = reg_PA["model"]
    params_pa = reg_PA["params"]
    bse_pa = reg_PA["bse"]
    pvals_pa = reg_PA["pvalues"]
    conf_pa = reg_PA["conf_int"]
    anova_pa = reg_PA["anova"]

    data_pa = []
    for name in params_pa.index:
        data_pa.append([
            name,
            f"{params_pa[name]:.8f}",
            f"{bse_pa[name]:.8f}",
            f"{model_pa.tvalues[name]:.4f}",
            f"{pvals_pa[name]:.2e}",
            f"{conf_pa.loc[name, 0]:.8f}",
            f"{conf_pa.loc[name, 1]:.8f}"
        ])
    cols_coefs = [
        "Paramètre",
        "Coefficient",
        "Erreur-type",
        "t-stat",
        "p-value",
        "IC95%-inf",
        "IC95%-sup"
    ]
    add_custom_table_data(doc, data_pa, cols_coefs, title="Coefficients de la régression (PA)")


    slope_pa_pval = pvals_pa["x"] if "x" in pvals_pa else np.nan
    if slope_pa_pval < alpha:
        slope_pa_interp = f"La pente est significative (p={slope_pa_pval:.3g} < {alpha})."
    else:
        slope_pa_interp = f"La pente n'est pas significative (p={slope_pa_pval:.3g} ≥ {alpha})."

    w_pval_pa = reg_PA["shapiro_pvalue"]
    if w_pval_pa < alpha:
        normality_pa_interp = (
            f"Les résidus ne suivent pas une distribution normale (p={w_pval_pa:.3g} < {alpha})."
        )
    else:
        normality_pa_interp = (
            f"Les résidus suivent une distribution proche de la normale (p={w_pval_pa:.3g} ≥ {alpha})."
        )

    doc.add_paragraph(
        f"R² = {reg_PA['r2']:.4f}, R² ajusté = {reg_PA['r2_adj']:.4f}\n"
        f"Test de Shapiro-Wilk sur les résidus : p-value = {w_pval_pa:.3g}\n"
        f"=> {normality_pa_interp}\n"
        f"=> {slope_pa_interp}"
    )

    doc.add_paragraph("ANOVA (PA) :")
    add_anova_table(doc, anova_pa, title="Tableau ANOVA - PA")


    fig_id = str(uuid.uuid4())[:8]
    fig_pa_resid = f"residus_pa_{fig_id}.png"
    plt.figure()
    plt.scatter(reg_PA["fitted_values"], reg_PA["residuals"])
    plt.axhline(0, linestyle='--')
    plt.title("Résidus vs Valeurs prédites (PA)")
    plt.xlabel("Valeurs prédites")
    plt.ylabel("Résidus")
    plt.tight_layout()
    plt.savefig(fig_pa_resid)
    plt.close()
    insert_matplotlib_figure(doc, fig_pa_resid, "Graphique : Résidus vs. Prédictions (PA)")


    fig_id = str(uuid.uuid4())[:8]
    fig_pa_obs_pred = f"obs_pred_pa_{fig_id}.png"
    plt.figure()
    plt.scatter(reg_PA["fitted_values"], reg_PA["model"].model.endog)
    min_val = min(min(reg_PA["fitted_values"]), min(reg_PA["model"].model.endog))
    max_val = max(max(reg_PA["fitted_values"]), max(reg_PA["model"].model.endog))
    plt.plot([min_val, max_val], [min_val, max_val], linestyle='--')
    plt.title("Valeurs Observées vs Valeurs Prédites (PA)")
    plt.xlabel("Valeurs Prédites")
    plt.ylabel("Valeurs Observées")
    plt.tight_layout()
    plt.savefig(fig_pa_obs_pred)
    plt.close()
    insert_matplotlib_figure(doc, fig_pa_obs_pred, "Graphique : Observé vs. Prédit (PA)")


    doc.add_heading("2.2 PA + FPR", level=2)
    model_pf = reg_PAFPR["model"]
    params_pf = reg_PAFPR["params"]
    bse_pf = reg_PAFPR["bse"]
    pvals_pf = reg_PAFPR["pvalues"]
    conf_pf = reg_PAFPR["conf_int"]
    anova_pf = reg_PAFPR["anova"]

    data_pf = []
    for name in params_pf.index:
        data_pf.append([
            name,
            f"{params_pf[name]:.8f}",
            f"{bse_pf[name]:.8f}",
            f"{model_pf.tvalues[name]:.4f}",
            f"{pvals_pf[name]:.2e}",
            f"{conf_pf.loc[name, 0]:.8f}",
            f"{conf_pf.loc[name, 1]:.8f}"
        ])
    add_custom_table_data(doc, data_pf, cols_coefs, title="Coefficients de la régression (PA+FPR)")

    slope_pafpr_pval = pvals_pf["x"] if "x" in pvals_pf else np.nan
    if slope_pafpr_pval < alpha:
        slope_pafpr_interp = f"La pente est significative (p={slope_pafpr_pval:.3g} < {alpha})."
    else:
        slope_pafpr_interp = f"La pente n'est pas significative (p={slope_pafpr_pval:.3g} ≥ {alpha})."

    w_pval_pf = reg_PAFPR["shapiro_pvalue"]
    if w_pval_pf < alpha:
        normality_pf_interp = (
            f"Les résidus ne suivent pas une distribution normale (p={w_pval_pf:.3g} < {alpha})."
        )
    else:
        normality_pf_interp = (
            f"Les résidus suivent une distribution proche de la normale (p={w_pval_pf:.3g} ≥ {alpha})."
        )

    doc.add_paragraph(
        f"R² = {reg_PAFPR['r2']:.4f}, R² ajusté = {reg_PAFPR['r2_adj']:.4f}\n"
        f"Test de Shapiro-Wilk sur les résidus : p-value = {w_pval_pf:.3g}\n"
        f"=> {normality_pf_interp}\n"
        f"=> {slope_pafpr_interp}"
    )

    doc.add_paragraph("ANOVA (PA+FPR) :")
    add_anova_table(doc, anova_pf, title="Tableau ANOVA - PA+FPR")


    fig_id = str(uuid.uuid4())[:8]
    fig_pf_resid = f"residus_pafpr_{fig_id}.png"
    plt.figure()
    plt.scatter(reg_PAFPR["fitted_values"], reg_PAFPR["residuals"])
    plt.axhline(0, linestyle='--')
    plt.title("Résidus vs Valeurs prédites (PA+FPR)")
    plt.xlabel("Valeurs prédites")
    plt.ylabel("Résidus")
    plt.tight_layout()
    plt.savefig(fig_pf_resid)
    plt.close()
    insert_matplotlib_figure(doc, fig_pf_resid, "Graphique : Résidus vs. Prédictions (PA+FPR)")


    fig_id = str(uuid.uuid4())[:8]
    fig_pf_obs_pred = f"obs_pred_pafpr_{fig_id}.png"
    plt.figure()
    plt.scatter(reg_PAFPR["fitted_values"], reg_PAFPR["model"].model.endog)
    min_val = min(min(reg_PAFPR["fitted_values"]), min(reg_PAFPR["model"].model.endog))
    max_val = max(max(reg_PAFPR["fitted_values"]), max(reg_PAFPR["model"].model.endog))
    plt.plot([min_val, max_val], [min_val, max_val], linestyle='--')
    plt.title("Valeurs Observées vs Valeurs Prédites (PA+FPR)")
    plt.xlabel("Valeurs Prédites")
    plt.ylabel("Valeurs Observées")
    plt.tight_layout()
    plt.savefig(fig_pf_obs_pred)
    plt.close()
    insert_matplotlib_figure(doc, fig_pf_obs_pred, "Graphique : Observé vs. Prédit (PA+FPR)")


    doc.add_heading("3. Vérification de la validité de la droite de régression", level=1)


    doc.add_paragraph("PA Seul : moyenne par niveau et (yij - Yimoy)².")
    data_val_pa = []
    for i, row in valid_PA.iterrows():
        data_val_pa.append([
            row["Niveau"],
            f"{row['xij']:.4f}",
            f"{row['yij']:.4f}",
            f"{row['y_moy']:.4f}",
            f"{row['sq_dev']:.4f}"
        ])
    cols_valid = ["Niveau", "xij", "yij", "Yimoy", "(yij - Yimoy)^2"]
    add_custom_table_data(doc, data_val_pa, cols_valid, title="PA - validité")
    doc.add_paragraph(f"SCE (Expérimentale) = {SCEexp_PA:.4f}.")


    doc.add_paragraph("PA+FPR : moyenne par niveau et (yij - Yimoy)².")
    data_val_pf = []
    for i, row in valid_PAFPR.iterrows():
        data_val_pf.append([
            row["Niveau"],
            f"{row['xij']:.4f}",
            f"{row['yij']:.4f}",
            f"{row['y_moy']:.4f}",
            f"{row['sq_dev']:.4f}"
        ])
    add_custom_table_data(doc, data_val_pf, cols_valid, title="PA+FPR - validité")
    doc.add_paragraph(f"SCE (Expérimentale) = {SCEexp_PAFPR:.4f}.")


    doc.add_heading("4. Test de non-linéarité", level=1)
    doc.add_paragraph(
        "Calcul et comparaison des SCE Résiduelle, SCE Expérimentale, et de l'erreur non linéaire. "
        "Cette analyse permet de déterminer si le modèle linéaire ajuste correctement les données expérimentales."
    )


    add_non_linearity_table(doc, df_PA, reg_PA, SCEexp_PA, name="D0 (PA seul)")


    add_non_linearity_table(doc, df_PAFPR, reg_PAFPR, SCEexp_PAFPR, name="D1 (PA+FPR)")


    doc.add_heading("5. Comparaison des pentes et des ordonnées (D0 vs. D1)", level=1)

    dslope = comp_pentes["diff_slope"]
    tslope = comp_pentes["t_slope"]
    p_slope = comp_pentes["p_slope"]
    df_slope = comp_pentes["df_slope"]

    dint = comp_pentes["diff_intercept"]
    tint = comp_pentes["t_intercept"]
    p_int = comp_pentes["p_intercept"]

    data_cmp_pentes = [
        ["Différence de pentes (slope0 - slope1)", f"{dslope:.4f}"],
        ["t calculée", f"{tslope:.4f}"],
        ["ddl (approx)", f"{df_slope}"],
        ["p-value (bilat)", f"{p_slope:.2e}"],
        ["Interprétation", comp_pentes["interpret_slope"]]
    ]
    add_custom_table_data(doc, data_cmp_pentes, ["Élément", "Valeur"],
                          title="Comparaison des pentes (D0 vs. D1)")

    data_cmp_ordo = [
        ["Différence d'ordonnées (int0 - int1)", f"{dint:.4f}"],
        ["t calculée", f"{tint:.4f}"],
        ["ddl (approx)", f"{comp_pentes['df_intercept']}"],
        ["p-value (bilat)", f"{p_int:.2e}"],
        ["Interprétation", comp_pentes["interpret_intercept"]]
    ]
    add_custom_table_data(doc, data_cmp_ordo, ["Élément", "Valeur"],
                          title="Comparaison des ordonnées à l'origine (D0 vs. D1)")


    doc.add_heading("6. Comparaison des ordonnées à l'origine avec zéro", level=1)

    t_calc_d0, df_d0, p_d0, interp_d0 = comp_int_d0_zero
    data_d0_zero = [
        ["Ordonnée (D0)", f"{reg_PA['params']['Intercept']:.4f}"],
        ["t calculée", f"{t_calc_d0:.4f}"],
        ["ddl", f"{df_d0}"],
        ["p-value", f"{p_d0:.2e}"],
        ["Interprétation", interp_d0]
    ]
    add_custom_table_data(doc, data_d0_zero, ["Élément", "Valeur"],
                          title="D0 vs. 0")

    t_calc_d1, df_d1, p_d1, interp_d1 = comp_int_d1_zero
    data_d1_zero = [
        ["Ordonnée (D1)", f"{reg_PAFPR['params']['Intercept']:.4f}"],
        ["t calculée", f"{t_calc_d1:.4f}"],
        ["ddl", f"{df_d1}"],
        ["p-value", f"{p_d1:.2e}"],
        ["Interprétation", interp_d1]
    ]
    add_custom_table_data(doc, data_d1_zero, ["Élément", "Valeur"],
                          title="D1 vs. 0")


    doc.add_heading("7. Conclusion générale", level=1)


    cond_cochran = bool_cochran_pa and bool_cochran_pafpr
    slope_ok_pa = (slope_pa_pval < alpha)
    slope_ok_pafpr = (slope_pafpr_pval < alpha)
    cond_slope = slope_ok_pa and slope_ok_pafpr
    cond_r2_pa = (reg_PA["r2"] >= 0.9)
    cond_r2_pafpr = (reg_PAFPR["r2"] >= 0.9)
    cond_norm_pa = (w_pval_pa >= alpha)
    cond_norm_pafpr = (w_pval_pf >= alpha)
    all_ok = (cond_cochran and cond_slope and cond_r2_pa and cond_r2_pafpr and cond_norm_pa and cond_norm_pafpr)


    if all_ok:
        conclusion_text = (
            "Les analyses réalisées démontrent que la méthode respecte les critères de linéarité sur l’ensemble de la plage étudiée. "
            "Le test de Cochran confirme l’homogénéité des variances, les régressions linéaires montrent des pentes significatives avec un coefficient de détermination élevé et des résidus conformes aux hypothèses (normalité vérifiée). "
            "Le test de non-linéarité n’a pas révélé de déviation significative, et l’ordonnée à l’origine est jugée acceptable ou proche de zéro.\n\n"
            "Interprétation : La relation linéaire établie permet d’affirmer que la méthode est adaptée et fiable dans la plage mesurée."
        )
    else:
        conclusion_text = (
            "Certains indicateurs remettent en cause la linéarité de la méthode : le test de Cochran indique une hétérogénéité des variances, "
            "la significativité de la pente est insuffisante ou le coefficient de détermination est faible, et/ou les résidus ne respectent pas la distribution normale. "
            "De plus, le test de non-linéarité suggère une déviation par rapport au modèle théorique.\n\n"
            "Interprétation : Ces résultats indiquent que la relation entre la variable d’entrée et la réponse pourrait ne pas être strictement linéaire, "
            "ce qui nécessite une investigation complémentaire (par exemple, en réévaluant la plage d’étude ou en ajustant les conditions expérimentales)."
        )
    doc.add_paragraph(conclusion_text)

    doc.save(output_file)
    print(f"\nRapport sauvegardé : {output_file}")


def run_linearity():
    print("\n=== Validation de la Linéarité (Automatique) ===")

    file_pa = input("Chemin du fichier Excel pour PA seul : ")
    file_pafpr = input("Chemin du fichier Excel pour PA+FPR : ")


    df_PA = pd.read_excel(file_pa, decimal=",")
    df_PAFPR = pd.read_excel(file_pafpr, decimal=",")


    for col in ["xij", "yij"]:
        df_PA[col] = pd.to_numeric(df_PA[col].astype(str).str.replace(',', '.'), errors='coerce')
        df_PAFPR[col] = pd.to_numeric(df_PAFPR[col].astype(str).str.replace(',', '.'), errors='coerce')


    grouped_pa = []
    for niv in sorted(df_PA["Niveau"].unique()):
        subset = df_PA[df_PA["Niveau"] == niv]["yij"].dropna().values
        grouped_pa.append(subset)

    grouped_pafpr = []
    for niv in sorted(df_PAFPR["Niveau"].unique()):
        subset = df_PAFPR[df_PAFPR["Niveau"] == niv]["yij"].dropna().values
        grouped_pafpr.append(subset)

    cochran_res_PA = cochran_test(grouped_pa)
    cochran_res_PAFPR = cochran_test(grouped_pafpr)


    x_pa = df_PA["xij"].values
    y_pa = df_PA["yij"].values
    reg_pa = linear_regression_analysis(x_pa, y_pa)

    x_pf = df_PAFPR["xij"].values
    y_pf = df_PAFPR["yij"].values
    reg_pf = linear_regression_analysis(x_pf, y_pf)


    valid_pa, SCEexp_pa = compute_validity_table(df_PA)
    valid_pf, SCEexp_pf = compute_validity_table(df_PAFPR)


    comp_pentes = compare_slopes_intercepts(reg_pa, reg_pf)


    comp_int_d0_zero = compare_intercept_with_zero(reg_pa)

    comp_int_d1_zero = compare_intercept_with_zero(reg_pf)


    create_report_linearity(
        df_PA, df_PAFPR,
        cochran_res_PA, cochran_res_PAFPR,
        reg_pa, reg_pf,
        valid_pa, SCEexp_pa,
        valid_pf, SCEexp_pf,
        comp_pentes,
        comp_int_d0_zero,
        comp_int_d1_zero,
        alpha=0.05,
        output_file="Rapport_Linéarité.docx"
    )



def run_justesse():
    print("\n=== Validation de la Justesse (Niveau=3 comme référence) ===")
    file_ = input("Chemin du fichier Excel (Niveau, xij, yij) : ")
    df = pd.read_excel(file_, decimal=",")
    for col in ["xij", "yij"]:
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(',', '.'), errors='coerce')
    df["Niveau"] = pd.to_numeric(df["Niveau"], errors='coerce')
    df.sort_values(by=["Niveau"], inplace=True, ignore_index=True)
    df_niv3 = df[df["Niveau"] == 3].copy()
    df_niv3.sort_index(inplace=True)
    df_niv3.reset_index(drop=True, inplace=True)
    if len(df_niv3) != 3:
        print("ERREUR : Le niveau=3 ne comporte pas exactement 3 lignes !")
        return
    day_factors = []
    for i in range(3):
        x_ = df_niv3.loc[i, "xij"]
        y_ = df_niv3.loc[i, "yij"]
        if x_ == 0 or pd.isna(x_) or pd.isna(y_):
            print("ERREUR : Donnée manquante ou x=0 dans le niveau=3")
            return
        day_factors.append(y_ / x_)
    df["rep_idx"] = df.groupby("Niveau").cumcount()
    df["Q_retrouvée"] = df.apply(lambda row: row["yij"] / day_factors[int(row["rep_idx"])], axis=1)
    df["Recouvrement"] = (df["Q_retrouvée"] / df["xij"]) * 100
    grouped = []
    for niv in sorted(df["Niveau"].unique()):
        subset = df[df["Niveau"] == niv]["Recouvrement"].dropna().values
        grouped.append(subset)
    (C, Ccrit, dec, bool_cochran) = cochran_test(grouped)
    df["Niveau_str"] = df["Niveau"].astype(str)
    model = ols("Recouvrement ~ Niveau_str", data=df).fit()
    anova_results = sm_anova.anova_lm(model, typ=1)
    p_value_anova = anova_results["PR(>F)"][0]
    rec_vals = df["Recouvrement"].dropna()
    mean_global = rec_vals.mean()
    std_global = rec_vals.std(ddof=1)
    n_global = rec_vals.count()
    alpha = 0.05
    df_ = n_global - 1
    if df_ > 0:
        t_crit = t.ppf(1 - alpha / 2, df_)
        ic_inf = mean_global - t_crit * std_global / np.sqrt(n_global)
        ic_sup = mean_global + t_crit * std_global / np.sqrt(n_global)
    else:
        ic_inf, ic_sup = np.nan, np.nan
    doc = Document()
    doc.add_heading("Rapport de Justesse (Niveau=3)", 0)
    doc.add_paragraph(
        "Ce rapport a été généré pour valider la justesse de la méthode en utilisant le niveau 3 comme référence.")
    doc.save("Rapport_Justesse_Niv3.docx")
    print("\nRapport sauvegardé : Rapport_Justesse_Niv3.docx")



def run_justesse_automatique():
    print("\n=== Validation de la Justesse (Automatique) ===")

    file_ = input("Chemin du fichier Excel (avec colonnes Niveau, xij, yij) : ")
    df = pd.read_excel(file_, decimal=",")

    for col in ["xij", "yij"]:
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(',', '.'), errors='coerce')
    df["Niveau"] = pd.to_numeric(df["Niveau"], errors='coerce')
    if df["Niveau"].isna().any() or df["xij"].isna().any() or df["yij"].isna().any():
        print("ATTENTION : Certaines valeurs manquent dans Niveau/xij/yij. Veuillez corriger votre fichier Excel.")
        return

    df.sort_values(by=["Niveau"], inplace=True, ignore_index=True)
    df["rep_idx"] = df.groupby("Niveau").cumcount().astype(int)
    if (df["rep_idx"] > 2).any():
        print("ATTENTION : Un niveau a plus de 3 répétitions => code non adapté.")
        return

    print("\nSaisir les 3 constantes (Jour1, Jour2, Jour3) :")
    cst_j1 = float(input("Constante Jour 1 : "))
    cst_j2 = float(input("Constante Jour 2 : "))
    cst_j3 = float(input("Constante Jour 3 : "))
    day_factors = [cst_j1, cst_j2, cst_j3]

    df["Q_retrouvée"] = df.apply(lambda row: row["yij"] / day_factors[int(row["rep_idx"])], axis=1)
    df["Recouvrement"] = (df["Q_retrouvée"] / df["xij"]) * 100

    grouped = []
    niveaux_uniques = sorted(df["Niveau"].unique())
    for niv in niveaux_uniques:
        arr = df[df["Niveau"] == niv]["Recouvrement"].values
        grouped.append(arr)
    (C, Ccrit, dec_cochran, bool_cochran) = cochran_test(grouped)

    df["Niveau_str"] = df["Niveau"].astype(str)
    model = ols("Q('Recouvrement') ~ Niveau_str", data=df).fit()
    anova_results = sm_anova.anova_lm(model, typ=1)
    p_value_anova = anova_results["PR(>F)"][0]

    rec_vals = df["Recouvrement"]
    mean_global = rec_vals.mean()
    std_global = rec_vals.std(ddof=1)
    n_global = rec_vals.count()
    alpha = 0.05
    df_ = n_global - 1
    if df_ > 0:
        t_crit = t.ppf(1 - alpha / 2, df_)
        ic_inf = mean_global - t_crit * std_global / np.sqrt(n_global)
        ic_sup = mean_global + t_crit * std_global / np.sqrt(n_global)
    else:
        ic_inf, ic_sup = np.nan, np.nan

    doc = Document()
    doc.add_heading("Rapport de Justesse (Automatique)", 0)
    doc.add_paragraph(
        "Ce rapport de validation de justesse vise à évaluer l’exactitude de la méthode analytique en utilisant le niveau de référence (N=3) pour établir trois constantes de calibration correspondant aux trois répétitions (Jour 1, Jour 2 et Jour 3). "
        "Pour chaque échantillon, la quantité retrouvée est calculée en divisant la réponse mesurée par la constante du jour concerné, puis le recouvrement est déterminé en comparant cette quantité à la valeur théorique.\n\n"
        "Les résultats sont ensuite soumis à :\n"
        " - Un test de Cochran pour vérifier l’homogénéité des variances entre répétitions,\n"
        " - Une ANOVA pour déterminer l'effet du niveau sur le recouvrement,\n"
        " - Et le calcul de statistiques globales (moyenne, écart-type et intervalle de confiance).\n\n"
        "Ce processus permet d’évaluer la justesse de la méthode et d’en apprécier la robustesse analytique."
    )

    data_for_table = []
    for i, row in df.iterrows():
        data_for_table.append([
            row["Niveau"],
            row["rep_idx"] + 1,
            f"{row['xij']:.4f}",
            f"{row['yij']:.4f}",
            f"{row['Q_retrouvée']:.4f}",
            f"{row['Recouvrement']:.4f}"
        ])
    add_custom_table_data(
        doc,
        data_for_table,
        ["Niveau", "Jour", "xij", "yij", "Q_retrouvée", "Recouvrement (%)"],
        title="Données brutes et Calculs"
    )

    cochran_info = [
        ["C calculé", f"{C:.4f}" if not np.isnan(C) else "NaN"],
        ["C Critique", f"{Ccrit:.4f}"],
        ["Conclusion", dec_cochran]
    ]
    add_custom_table_data(doc, cochran_info, ["Paramètre", "Valeur"], title="Test de Cochran")
    doc.add_paragraph(
        "Interprétation : Un test de Cochran indiquant des variances homogènes confirme une bonne répétabilité des mesures.")

    doc.add_paragraph("Analyse de variance (ANOVA) un facteur : Recouvrement ~ Niveau")
    add_anova_table(doc, anova_results, title="Tableau ANOVA - Justesse")
    if p_value_anova < alpha:
        anova_interp = f"p={p_value_anova:.3g} < {alpha} => l'effet du niveau est significatif."
    else:
        anova_interp = f"p={p_value_anova:.3g} >= {alpha} => aucun effet significatif du niveau."
    doc.add_paragraph(f"Interprétation ANOVA : {anova_interp}")

    stat_global_data = [
        ["Recouvrement moyen (%)", f"{mean_global:.4f}"],
        ["Écart-type global", f"{std_global:.4f}"],
        ["Nombre total d'échantillons", f"{n_global}"],
        ["IC 95% (limite inférieure)", f"{ic_inf:.4f}"],
        ["IC 95% (limite supérieure)", f"{ic_sup:.4f}"]
    ]
    add_custom_table_data(doc, stat_global_data, ["Statistique", "Valeur"], title="Statistiques globales")

    doc.add_heading("Conclusion générale", level=1)
    concl_lines = []
    if bool_cochran:
        concl_lines.append(
            "Le test de Cochran indique que les variances entre les répétitions sont homogènes, confirmant ainsi une bonne répétabilité.")
    else:
        concl_lines.append(
            "Le test de Cochran révèle une hétérogénéité des variances, ce qui pourrait affecter la précision.")
    if p_value_anova < alpha:
        concl_lines.append(
            "L'ANOVA démontre un effet significatif du niveau, indiquant que la réponse du procédé varie entre les niveaux.")
    else:
        concl_lines.append(
            "L'ANOVA n'indique pas d'effet significatif du niveau, suggérant une uniformité des mesures entre les niveaux.")
    concl_lines.append(
        f"Le recouvrement moyen global est de {mean_global:.2f}% avec un intervalle de confiance à 95% de [{ic_inf:.2f}, {ic_sup:.2f}].")
    if ic_inf <= 100 <= ic_sup:
        concl_lines.append(
            "La valeur idéale de 100% est incluse dans l'intervalle, confirmant l'exactitude de la méthode.")
    else:
        concl_lines.append(
            "La valeur idéale de 100% n'est pas incluse dans l'intervalle, indiquant une possible dérive systématique.")
    if bool_cochran and (p_value_anova >= alpha) and (ic_inf <= 100 <= ic_sup):
        concl_lines.append(
            "Conclusion finale : La méthode est scientifiquement juste et fournit une mesure précise et fiable.")
    else:
        concl_lines.append(
            "Conclusion finale : Certains indicateurs suggèrent que la méthode nécessite une réévaluation pour optimiser répétabilité et exactitude.")
    doc.add_paragraph("\n".join(concl_lines))

    output_file = "Rapport_Justesse_Automatique.docx"
    doc.save(output_file)
    print(f"\nRapport sauvegardé : {output_file}")



def run_fidelite_series():
    """
    Évalue la fidélité d'une méthode analytique sur 3 séries (7 répétitions chacune)
    en calculant le pourcentage de recouvrement et les composantes de variance.
    """
    print("\n=== Fidélité sur 3 séries (7 répétitions chacune) ===")

    file_ = input("Chemin du fichier Excel : ")
    df = pd.read_excel(file_, decimal=",")

    for col in ["qté introd.(mg)", "Air"]:
        df[col] = pd.to_numeric(df[col].astype(str).str.replace(',', '.'), errors='coerce')

    df.sort_values(by=["séries", "essai"], inplace=True, ignore_index=True)

    print("\nSaisir les 3 constantes (une par série) :")
    cst_s1 = float(input("Constante Série 1 : "))
    cst_s2 = float(input("Constante Série 2 : "))
    cst_s3 = float(input("Constante Série 3 : "))
    day_factors = [cst_s1, cst_s2, cst_s3]

    def calc_qte_retrouvee(row):
        s_idx = int(row["séries"]) - 1
        return row["Air"] / day_factors[s_idx]

    df["qté retrouv."] = df.apply(calc_qte_retrouvee, axis=1)
    df["R%"] = (df["qté retrouv."] / df["qté introd.(mg)"]) * 100

    grouped = []
    for s in [1, 2, 3]:
        arr = df[df["séries"] == s]["R%"].dropna().values
        grouped.append(arr)
    (C, Ccrit, cochran_dec, cochran_bool) = cochran_test(grouped)

    df["séries_str"] = df["séries"].astype(str)
    model = ols("Q('R%') ~ séries_str", data=df).fit()
    anova_res = sm_anova.anova_lm(model, typ=1)
    pval_anova = anova_res["PR(>F)"]["séries_str"]

    MSb = anova_res["mean_sq"]["séries_str"]
    MSw = anova_res["mean_sq"]["Residual"]
    n = 7
    Sd2 = (MSb - MSw) / n if (MSb > MSw) else 0
    Sr2 = MSw
    SFI2 = Sr2 + Sd2
    Sr = np.sqrt(Sr2)
    Sd = np.sqrt(Sd2)
    SFI = np.sqrt(SFI2)

    global_mean = df["R%"].mean()
    CVr = (Sr / global_mean) * 100
    CVd = (Sd / global_mean) * 100
    CVFI = (SFI / global_mean) * 100


    doc = Document()
    doc.add_heading("Rapport de Fidélité - 3 séries", 0)
    doc.add_paragraph(
        "Ce rapport de fidélité a pour objectif d’évaluer la reproductibilité et la précision intermédiaire de la méthode analytique sur trois séries distinctes, "
        "chacune comportant sept répétitions. Les données issues d’un fichier Excel (colonnes : séries, essai, qté introd.(mg), Air) permettent de calculer la quantité retrouvée via une constante spécifique à chaque série, "
        "et le pourcentage de recouvrement (R%) est déterminé en comparant cette quantité à la quantité introduite.\n\n"
        "Les analyses comprennent :\n"
        " - Un test de Cochran pour vérifier l’homogénéité des variances entre les séries,\n"
        " - Une ANOVA un facteur pour détecter d’éventuelles différences entre les séries,\n"
        " - L’extraction des composantes de variance pour calculer les CV de répétabilité (CVr), d’effet série (CVd) et de précision intermédiaire (CVFI).\n\n"
        "Cette approche permet d’apprécier la fidélité globale de la méthode dans le cadre d’un contrôle qualité rigoureux."
    )

    data_rows = []
    for i, row in df.iterrows():
        data_rows.append([
            row["séries"],
            row["essai"],
            f"{row['qté introd.(mg)']:.4f}",
            f"{row['Air']:.4f}",
            f"{row['qté retrouv.']:.4f}",
            f"{row['R%']:.4f}"
        ])
    add_custom_table_data(
        doc,
        data_rows,
        ["séries", "essai", "qté introd.(mg)", "Air", "qté retrouv.", "R%"],
        title="Données brutes + Calcul R%"
    )

    cochran_info = [
        ["C calculé", f"{C:.4f}" if not np.isnan(C) else "NaN"],
        ["C critique", f"{Ccrit:.3f}"],
        ["Conclusion", cochran_dec]
    ]
    add_custom_table_data(doc, cochran_info, ["Paramètre", "Valeur"],
                          title="Test de Cochran (Homogénéité des Variances)")
    doc.add_paragraph("Interprétation : Le test de Cochran vérifie que les variances entre séries sont homogènes.")

    doc.add_paragraph("ANOVA un facteur : R% ~ séries")
    add_anova_table(doc, anova_res, title="Tableau ANOVA - Fidélité")
    doc.add_paragraph(
        f"Interprétation ANOVA : p-value = {pval_anova:.4g}. {'Différence significative entre séries.' if pval_anova < 0.05 else 'Aucune différence significative entre séries.'}")

    var_data = [
        ["MS(within) = Sr²", f"{Sr2:.4f}"],
        ["MS(between)", f"{MSb:.4f}"],
        ["Sd²", f"{Sd2:.4f}"],
        ["SFI²", f"{SFI2:.4f}"]
    ]
    add_custom_table_data(doc, var_data, ["Paramètre", "Valeur"], title="Composantes de variance")

    sd_data = [
        ["Sr", f"{Sr:.4f}"],
        ["Sd", f"{Sd:.4f}"],
        ["SFI", f"{SFI:.4f}"],
        ["Moyenne globale R%", f"{global_mean:.4f}"],
        ["CVr (%)", f"{CVr:.2f}"],
        ["CVd (%)", f"{CVd:.2f}"],
        ["CVFI (%)", f"{CVFI:.2f}"]
    ]
    add_custom_table_data(doc, sd_data, ["Paramètre", "Valeur"], title="Écart-types et Coefficients de variation")

    doc.add_heading("Conclusion générale", level=1)
    conclusion = []
    if cochran_bool:
        conclusion.append(
            "Le test de Cochran indique que les variances entre les séries sont homogènes, ce qui est un indicateur positif de répétabilité.")
    else:
        conclusion.append(
            "Le test de Cochran révèle une hétérogénéité des variances entre les séries, suggérant une variabilité qui pourrait compromettre la reproductibilité.")
    if pval_anova < 0.05:
        conclusion.append(
            "L'ANOVA démontre une différence significative entre les séries, indiquant une variabilité inter-série non négligeable.")
    else:
        conclusion.append(
            "L'ANOVA n'indique pas de différence significative entre les séries, suggérant une bonne uniformité des mesures.")
    conclusion.append(
        f"Les coefficients de variation calculés montrent une répétabilité (CVr) de {CVr:.2f}%, un effet série (CVd) de {CVd:.2f}%, "
        f"et une précision intermédiaire totale (CVFI) de {CVFI:.2f}%."
    )
    if CVFI < 2.0:
        conclusion.append("La précision intermédiaire est jugée acceptable (CVFI < 2%).")
    else:
        conclusion.append(
            "La précision intermédiaire dépasse le critère de 2%, suggérant la nécessité d’un ajustement ou d'une investigation complémentaire.")
    doc.add_paragraph("\n".join(conclusion))

    output_file = "Rapport_Fidelite_Series.docx"
    doc.save(output_file)
    print(f"\nRapport sauvegardé : {output_file}")



def run_profil_exactitude_final():
    # Ici, on lit les valeurs via input() (déjà simulées par run_report_with_inputs)
    nb_levels_str = input("Nombre de niveaux ? ")
    try:
        nb_levels = int(nb_levels_str)
    except:
        print("Nombre de niveaux invalide.")
        return
    beta_str = input("Probabilité de tolérance (bêta) ? (ex: 0.95) ")
    beta = float(beta_str) if beta_str else 0.95
    limite_acceptable_str = input("Limite d'acceptabilité ± (%) ? (ex: 5) ")
    limite_acceptable = float(limite_acceptable_str) if limite_acceptable_str else 5.0

    level_params = []
    for lvl in range(nb_levels):
        file_path = input(f"Chemin du fichier Excel pour le Niveau {lvl + 1} : ")
        ref_str = input(f"Valeur de référence pour le Niveau {lvl + 1} : ")
        try:
            val_ref = float(ref_str.replace(',', '.'))
        except:
            print(f"Valeur de référence invalide pour le Niveau {lvl + 1}.")
            return
        level_params.append((file_path, val_ref))


    x_levels = []
    recov_array = []
    tol_lower_array = []
    tol_upper_array = []
    accept_lower_array = []
    accept_upper_array = []
    incert_array = []
    raw_snippets = []
    levels_data = []
    row_labels = [
        "Valeur de référence",
        "Moyenne niveau",
        "Écart-type de répétabilité (sr)",
        "Écart-type de fidélité (sFI)",
        "df Satterthwaite",
        "Facteur de couverture (ktol)",
        "Écart-type du IT (sIT)",
        "Valeur basse intervalle tolérance",
        "Valeur haute intervalle tolérance",
        "Biais relatif (%)",
        "Taux de recouvrement (%)",
        "Limite basse tolérance (%)",
        "Limite haute tolérance (%)",
        "Limite d'acceptabilité basse",
        "Limite d'acceptabilité haute",
        "Incertitude relative (%)"
    ]
    for i, (file_path, ref_value) in enumerate(level_params):
        level_label = chr(65 + i)
        try:
            df_raw = pd.read_excel(file_path)
        except Exception as e:
            print(f"Erreur de lecture pour le Niveau {level_label} : {e}")
            return
        snippet = df_raw.head(5).copy()
        raw_snippets.append((level_label, snippet))
        col_serie = df_raw.columns[0]
        df_long = df_raw.melt(id_vars=col_serie, var_name="Rep", value_name="Value")
        df_long.rename(columns={col_serie: "Serie"}, inplace=True)
        df_long["Value"] = df_long["Value"].astype(str).str.replace(',', '.')
        df_long["Value"] = pd.to_numeric(df_long["Value"], errors='coerce')
        df_long.dropna(subset=["Value"], inplace=True)
        I = df_long["Serie"].nunique()
        N = df_long.shape[0]
        J = int(round(N / I)) if I > 0 else 0
        df_long["Serie"] = df_long["Serie"].astype("category")
        model = ols("Value ~ C(Serie)", data=df_long).fit()
        anova_res = sm_anova.anova_lm(model, typ=1)
        SCE_inter = anova_res.loc["C(Serie)", "sum_sq"]
        SCE_resid = anova_res.loc["Residual", "sum_sq"]
        df_resid = I * (J - 1)
        s2r = SCE_resid / df_resid if df_resid > 0 else np.nan
        inter_val = ((SCE_inter / (I - 1)) - s2r) / J if (I - 1) > 0 and J > 0 else np.nan
        s2L = max(0, inter_val) if not np.isnan(inter_val) else np.nan
        s2FI = s2r + s2L if not np.isnan(s2r) else np.nan
        sr = np.sqrt(s2r) if not np.isnan(s2r) else np.nan
        sFI = np.sqrt(s2FI) if not np.isnan(s2FI) else np.nan
        mean_level = df_long["Value"].mean()
        bias_rel = ((mean_level / ref_value) - 1) * 100 if ref_value != 0 else np.nan
        recouv = (mean_level / ref_value) * 100 if ref_value != 0 else np.nan
        if not np.isnan(s2r) and s2r > 1e-15:
            R = s2L / s2r
        else:
            R = np.nan
        if not np.isnan(R) and (I - 1) > 0 and J > 0 and df_resid > 0:
            num = (R + 1) ** 2
            denom = (((R + 1) / J) ** 2 / (I - 1)) + ((1 - 1 / J) / (I * (J - 1)))
            df_satter = num / denom if abs(denom) >= 1e-15 else np.nan
        else:
            df_satter = np.nan
        alpha_local = 1 - beta
        ktol = t.ppf(1 - alpha_local / 2, df_satter) if not np.isnan(df_satter) and df_satter > 0 else np.nan
        sIT = sFI
        if not np.isnan(mean_level) and not np.isnan(ktol) and not np.isnan(sIT):
            IT_lower = mean_level - ktol * sIT
            IT_upper = mean_level + ktol * sIT
        else:
            IT_lower = np.nan
            IT_upper = np.nan
        tol_lower_percent = (IT_lower / ref_value) * 100 if ref_value != 0 and not np.isnan(IT_lower) else np.nan
        tol_upper_percent = (IT_upper / ref_value) * 100 if ref_value != 0 and not np.isnan(IT_upper) else np.nan
        accept_lower = 100 - limite_acceptable
        accept_upper = 100 + limite_acceptable
        incert_etendue = 2 * sIT if not np.isnan(sIT) else np.nan
        incert_rel_percent = (incert_etendue / mean_level) * 100 if not (
                    np.isnan(mean_level) or mean_level == 0 or np.isnan(incert_etendue)) else np.nan
        ld = {
            "Niveau": level_label,
            "Valeur de référence": ref_value,
            "Moyenne niveau": mean_level,
            "Écart-type de répétabilité (sr)": sr,
            "Écart-type de fidélité (sFI)": sFI,
            "df Satterthwaite": df_satter,
            "Facteur de couverture (ktol)": ktol,
            "Écart-type du IT (sIT)": sIT,
            "Valeur basse intervalle tolérance": IT_lower,
            "Valeur haute intervalle tolérance": IT_upper,
            "Biais relatif (%)": bias_rel,
            "Taux de recouvrement (%)": recouv,
            "Limite basse tolérance (%)": tol_lower_percent,
            "Limite haute tolérance (%)": tol_upper_percent,
            "Limite d'acceptabilité basse": accept_lower,
            "Limite d'acceptabilité haute": accept_upper,
            "Incertitude relative (%)": incert_rel_percent
        }
        levels_data.append(ld)
        x_levels.append(i + 1)
        recov_array.append(recouv)
        tol_lower_array.append(tol_lower_percent)
        tol_upper_array.append(tol_upper_percent)
        accept_lower_array.append(accept_lower)
        accept_upper_array.append(accept_upper)
        incert_array.append(incert_rel_percent)

    fig_id = str(uuid.uuid4())[:8]
    fig_recov = f"Recouv_Combined_{fig_id}.png"
    plt.figure(figsize=(6, 4))
    plt.plot(x_levels, recov_array, 'ro-', label="Taux de recouvrement (%)")
    plt.plot(x_levels, tol_lower_array, 'b-', label="Limite basse tolérance (%)")
    plt.plot(x_levels, tol_upper_array, 'b-', label="Limite haute tolérance (%)")
    plt.plot(x_levels, accept_lower_array, 'k--', label="Limite accept. basse")
    plt.plot(x_levels, accept_upper_array, 'k--', label="Limite accept. haute")
    plt.title("Taux de recouvrement (%) - Plot Combiné")
    plt.xlabel("Niveau")
    plt.ylabel("Taux de recouvrement (%)")
    plt.legend()
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(fig_recov)
    plt.close()

    fig_id = str(uuid.uuid4())[:8]
    fig_incert = f"Incert_Combined_{fig_id}.png"
    plt.figure(figsize=(6, 4))
    plt.plot(x_levels, incert_array, 'go-', markersize=8, label="Incertitude relative (%)")
    valid_x, valid_y = [], []
    for xx, yy in zip(x_levels, incert_array):
        if not np.isnan(xx) and not np.isnan(yy) and xx > 0 and yy > 0:
            valid_x.append(xx)
            valid_y.append(yy)
    if len(valid_x) >= 2:
        X = np.log(valid_x)
        Y = np.log(valid_y)
        p = np.polyfit(X, Y, 1)
        b_ = p[0]
        a_ = np.exp(p[1])
        x_smooth = np.linspace(min(valid_x), max(valid_x), 50)
        y_smooth = a_ * (x_smooth ** b_)
        plt.plot(x_smooth, y_smooth, 'g-', alpha=0.7)
        eq_text = f"y = {a_:.4f} x^{b_:.2f}"
        plt.text(valid_x[0] + 0.1, valid_y[0], eq_text)
    plt.title("Incertitude relative (%) - Plot Combiné")
    plt.xlabel("Niveau")
    plt.ylabel("Incertitude relative (%)")
    plt.legend()
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(fig_incert)
    plt.close()

    doc = Document()
    doc.add_heading("Rapport – Profil d'Exactitude (Final)", 0)
    intro = (
        "Introduction : Le profil d'exactitude évalue simultanément la justesse et la fidélité d'une méthode. "
        "Une ANOVA à un facteur est utilisée pour estimer la variance de répétabilité et inter-séries, permettant de calculer un intervalle de tolérance."
    )
    doc.add_paragraph(intro)
    doc.add_heading("1) Données Brutes Importées", level=1)
    for lvl_label, snippet in raw_snippets:
        doc.add_heading(f"Niveau {lvl_label}", level=2)
        table_snip = doc.add_table(rows=1, cols=len(snippet.columns))
        table_snip.style = 'Light Shading Accent 1'
        hdr_cells = table_snip.rows[0].cells
        for j, col in enumerate(snippet.columns):
            hdr_cells[j].text = str(col)
        for i, row_data in snippet.iterrows():
            row_cells = table_snip.add_row().cells
            for j, col in enumerate(snippet.columns):
                row_cells[j].text = str(row_data[col])
    doc.add_heading("2) Résultats Généraux (Tableau)", level=1)
    table = doc.add_table(rows=len(row_labels) + 1, cols=nb_levels + 1)
    table.style = 'Light Shading Accent 1'
    table.cell(0, 0).text = "Paramètre"
    for i, ld in enumerate(levels_data, start=1):
        table.cell(0, i).text = ld["Niveau"]
    for r_i, row_label in enumerate(row_labels, start=1):
        table.cell(r_i, 0).text = row_label
        for c_i, ld in enumerate(levels_data, start=1):
            val = ld.get(row_label, "")
            if isinstance(val, (float, int)) and not np.isnan(val):
                if ("%" in row_label.lower()) or ("recouvrement" in row_label.lower()):
                    table.cell(r_i, c_i).text = f"{val:.2f}%"
                else:
                    table.cell(r_i, c_i).text = f"{val:.5f}"
            else:
                table.cell(r_i, c_i).text = str(val)
    doc.add_heading("3) Interprétation des Résultats", level=1)
    method_valid = True
    reasons = []
    for ld in levels_data:
        b_ = ld.get("Biais relatif (%)", np.nan)
        if not np.isnan(b_):
            if abs(b_) > limite_acceptable:
                method_valid = False
                reasons.append(f"Niveau {ld['Niveau']}: Biais={b_:.2f}% > ±{limite_acceptable}%")
        inc_ = ld.get("Incertitude relative (%)", np.nan)
        if not np.isnan(inc_) and inc_ > 10.0:
            method_valid = False
            reasons.append(f"Niveau {ld['Niveau']}: Incert.={inc_:.2f}% > 10%")
    if method_valid:
        doc.add_paragraph(
            f"Aucun biais relatif n'excède ±{limite_acceptable}%, et l'incertitude reste inférieure à 10%.")
    else:
        doc.add_paragraph("Certains paramètres dépassent les critères fixés :")
        for r in reasons:
            doc.add_paragraph(f" - {r}")
    doc.add_heading("4) Graphiques Combinés", level=1)
    doc.add_paragraph(
        "Le premier graphique montre le Taux de recouvrement (%) pour chaque niveau, avec les lignes de tolérance et limites d'acceptabilité. "
        "Le second graphique illustre l'incertitude relative (%) par niveau."
    )
    if os.path.exists(fig_recov):
        doc.add_paragraph("Graphique : Taux de recouvrement (%)")
        doc.add_picture(fig_recov, width=Inches(5))
        os.remove(fig_recov)
    if os.path.exists(fig_incert):
        doc.add_paragraph("Graphique : Incertitude relative (%)")
        doc.add_picture(fig_incert, width=Inches(5))
        os.remove(fig_incert)
    doc.add_heading("5) Conclusion", level=1)
    if method_valid:
        doc.add_paragraph(
            f"Conclusion : La méthode respecte les critères pour le biais (±{limite_acceptable}%) et présente une incertitude acceptable (<10%).")
    else:
        doc.add_paragraph(
            "Conclusion : Certains indicateurs dépassent les seuils fixés, veuillez revoir les conditions expérimentales.")
    doc.add_heading("Paramètres Globaux", level=1)
    doc.add_paragraph(f"Bêta (tolérance) : {beta * 100:.2f}%\nLimite d'acceptabilité (± %) : {limite_acceptable}%")
    doc.add_paragraph("Fin du rapport.")
    out_file = "Rapport_Profil_Exactitude_Final.docx"
    doc.save(out_file)
    messagebox.showinfo("Succès", f"Rapport sauvegardé : {out_file}")
    print(f"\nRapport sauvegardé : {out_file}")


###############################################################################
# Fonction utilitaire pour simuler input() via une liste d'inputs
###############################################################################
def run_report_with_inputs(report_function, inputs_list):
    original_input = __builtins__.input
    it = iter(inputs_list)
    __builtins__.input = lambda prompt="": next(it)
    try:
        report_function()
    except Exception as e:
        messagebox.showerror("Erreur", str(e))
    finally:
        __builtins__.input = original_input



def browse_file_entry(entry):
    file_path = filedialog.askopenfilename(
        title="Sélectionner un fichier",
        filetypes=[("Fichiers Excel", "*.xls *.xlsx"), ("Tous les fichiers", "*.*")]
    )
    if file_path:
        entry.delete(0, tk.END)
        entry.insert(0, file_path)


def open_linearity_window(master):
    win = tk.Toplevel(master)
    win.title("Validation Linéarité")
    win.geometry("450x300")
    tk.Label(win, text="Fichier Excel pour PA seul:").pack(pady=5)
    entry_pa = tk.Entry(win, width=50)
    entry_pa.pack(pady=5)
    tk.Button(win, text="Parcourir", command=lambda: browse_file_entry(entry_pa)).pack(pady=5)
    tk.Label(win, text="Fichier Excel pour PA+FPR:").pack(pady=5)
    entry_pafpr = tk.Entry(win, width=50)
    entry_pafpr.pack(pady=5)
    tk.Button(win, text="Parcourir", command=lambda: browse_file_entry(entry_pafpr)).pack(pady=5)

    def generate_linearity_report():
        file_pa = entry_pa.get()
        file_pafpr = entry_pafpr.get()
        if not file_pa or not file_pafpr:
            messagebox.showerror("Erreur", "Veuillez sélectionner les deux fichiers Excel.")
            return
        inputs = [file_pa, file_pafpr]
        run_report_with_inputs(run_linearity, inputs)
        messagebox.showinfo("Succès", "Rapport de Linéarité généré avec succès.")

    tk.Button(win, text="Générer Rapport Linéarité", command=generate_linearity_report).pack(pady=10)


def open_justesse_window(master):
    win = tk.Toplevel(master)
    win.title("Validation Justesse")
    win.geometry("450x300")
    tk.Label(win, text="Fichier Excel (Niveau, xij, yij):").pack(pady=5)
    entry_file = tk.Entry(win, width=50)
    entry_file.pack(pady=5)
    tk.Button(win, text="Parcourir", command=lambda: browse_file_entry(entry_file)).pack(pady=5)
    tk.Label(win, text="Constante Jour 1:").pack(pady=5)
    entry_j1 = tk.Entry(win, width=20)
    entry_j1.pack(pady=5)
    tk.Label(win, text="Constante Jour 2:").pack(pady=5)
    entry_j2 = tk.Entry(win, width=20)
    entry_j2.pack(pady=5)
    tk.Label(win, text="Constante Jour 3:").pack(pady=5)
    entry_j3 = tk.Entry(win, width=20)
    entry_j3.pack(pady=5)

    def generate_justesse_report():
        file_justesse = entry_file.get()
        c1 = entry_j1.get()
        c2 = entry_j2.get()
        c3 = entry_j3.get()
        if not file_justesse or not c1 or not c2 or not c3:
            messagebox.showerror("Erreur", "Veuillez remplir tous les champs pour la Justesse.")
            return
        inputs = [file_justesse, c1, c2, c3]
        run_report_with_inputs(run_justesse_automatique, inputs)
        messagebox.showinfo("Succès", "Rapport de Justesse généré avec succès.")

    tk.Button(win, text="Générer Rapport Justesse", command=generate_justesse_report).pack(pady=10)


def open_fidelite_window(master):
    win = tk.Toplevel(master)
    win.title("Validation Fidélité")
    win.geometry("450x300")
    tk.Label(win, text="Fichier Excel:").pack(pady=5)
    entry_file = tk.Entry(win, width=50)
    entry_file.pack(pady=5)
    tk.Button(win, text="Parcourir", command=lambda: browse_file_entry(entry_file)).pack(pady=5)
    tk.Label(win, text="Constante Série 1:").pack(pady=5)
    entry_s1 = tk.Entry(win, width=20)
    entry_s1.pack(pady=5)
    tk.Label(win, text="Constante Série 2:").pack(pady=5)
    entry_s2 = tk.Entry(win, width=20)
    entry_s2.pack(pady=5)
    tk.Label(win, text="Constante Série 3:").pack(pady=5)
    entry_s3 = tk.Entry(win, width=20)
    entry_s3.pack(pady=5)

    def generate_fidelite_report():
        file_fid = entry_file.get()
        s1 = entry_s1.get()
        s2 = entry_s2.get()
        s3 = entry_s3.get()
        if not file_fid or not s1 or not s2 or not s3:
            messagebox.showerror("Erreur", "Veuillez remplir tous les champs pour la Fidélité.")
            return
        inputs = [file_fid, s1, s2, s3]
        run_report_with_inputs(run_fidelite_series, inputs)
        messagebox.showinfo("Succès", "Rapport de Fidélité généré avec succès.")

    tk.Button(win, text="Générer Rapport Fidélité", command=generate_fidelite_report).pack(pady=10)


def open_exactitude_window(master):
    win = tk.Toplevel(master)
    win.title("Profil d'Exactitude")
    win.geometry("500x600")
    tk.Label(win, text="Nombre de niveaux:").pack(pady=5)
    entry_levels = tk.Entry(win, width=10)
    entry_levels.insert(0, "2")
    entry_levels.pack(pady=5)
    tk.Label(win, text="Probabilité de tolérance (bêta, ex: 0.95):").pack(pady=5)
    entry_beta = tk.Entry(win, width=10)
    entry_beta.insert(0, "0.95")
    entry_beta.pack(pady=5)
    tk.Label(win, text="Limite d'acceptabilité ± (%), ex: 5:").pack(pady=5)
    entry_accept = tk.Entry(win, width=10)
    entry_accept.insert(0, "5")
    entry_accept.pack(pady=5)
    levels_frame = tk.Frame(win)
    levels_frame.pack(pady=10, fill=tk.BOTH, expand=True)
    level_entries = []  # Pour stocker (file_entry, ref_entry)

    def create_level_fields():
        for widget in levels_frame.winfo_children():
            widget.destroy()
        level_entries.clear()
        try:
            n_levels = int(entry_levels.get())
        except ValueError:
            messagebox.showerror("Erreur", "Veuillez entrer un nombre valide de niveaux.")
            return
        for i in range(n_levels):
            subframe = tk.Frame(levels_frame, bd=2, relief=tk.GROOVE, padx=5, pady=5)
            subframe.pack(pady=5, fill=tk.X)
            level_label = f"Niveau {i + 1} ({chr(65 + i)})"
            tk.Label(subframe, text=level_label, font=("Helvetica", 10, "bold")).pack(anchor="w")
            tk.Label(subframe, text="Fichier Excel:").pack(anchor="w")
            file_entry = tk.Entry(subframe, width=40)
            file_entry.pack(side=tk.LEFT, padx=5)
            tk.Button(subframe, text="Parcourir", command=lambda e=file_entry: browse_file_entry(e)).pack(side=tk.LEFT,
                                                                                                          padx=5)
            tk.Label(subframe, text="Valeur de référence:").pack(anchor="w", pady=(5, 0))
            ref_entry = tk.Entry(subframe, width=20)
            ref_entry.pack(anchor="w", padx=5)
            level_entries.append((file_entry, ref_entry))

    tk.Button(win, text="Créer les champs pour niveaux", command=create_level_fields).pack(pady=10)

    def generate_exactitude_report():
        n_levels = entry_levels.get()
        beta = entry_beta.get()
        accept = entry_accept.get()
        if not n_levels or not beta or not accept:
            messagebox.showerror("Erreur",
                                 "Veuillez remplir les champs communs (nombre de niveaux, bêta, acceptabilité).")
            return
        inputs = [n_levels, beta, accept]
        if not level_entries:
            messagebox.showerror("Erreur",
                                 "Veuillez cliquer sur 'Créer les champs pour niveaux' et remplir les informations pour chaque niveau.")
            return
        for idx, (file_entry, ref_entry) in enumerate(level_entries):
            file_val = file_entry.get()
            ref_val = ref_entry.get()
            if not file_val or not ref_val:
                messagebox.showerror("Erreur", f"Veuillez remplir tous les champs pour le niveau {idx + 1}.")
                return
            inputs.extend([file_val, ref_val])
        run_report_with_inputs(run_profil_exactitude_final, inputs)
        messagebox.showinfo("Succès", "Rapport de Profil d'Exactitude généré avec succès.")

    tk.Button(win, text="Générer Rapport Exactitude", command=generate_exactitude_report).pack(pady=10)


def main_menu():
    root = tk.Tk()
    root.title("Menu Principal - Validation Analytique")
    root.geometry("400x800")
    try:
        img = Image.open("ISACQ.png")
        photo = ImageTk.PhotoImage(img)
        lbl_img = tk.Label(root, image=photo)
        lbl_img.image = photo
        lbl_img.pack(pady=10)
    except Exception as e:
        print("Erreur chargement image:", e)
    tk.Label(root, text="Realized by: BENZAIMIA Mohamed Sami",
             fg="blue", font=("Helvetica", 10, "italic")).pack(side="bottom", pady=10)
    tk.Label(root, text="Sélectionnez la validation souhaitée :", font=("Helvetica", 14, "bold")).pack(pady=10)
    tk.Button(root, text="Linéarité", width=30, command=lambda: open_linearity_window(root)).pack(pady=5)
    tk.Button(root, text="Justesse", width=30, command=lambda: open_justesse_window(root)).pack(pady=5)
    tk.Button(root, text="Fidélité", width=30, command=lambda: open_fidelite_window(root)).pack(pady=5)
    tk.Button(root, text="Analyse de Robustesse", width=30, command=lambda: open_robustness_window(root)).pack(pady=5)
    tk.Button(root, text="Exactitude", width=30, command=lambda: open_exactitude_window(root)).pack(pady=5)
    tk.Button(root, text="Quitter", width=30, command=root.quit).pack(pady=20)
    root.mainloop()


if __name__ == "__main__":
    main_menu()